"""
SGUI by I&K is an image steganography desktop application that uses the modified least significant bit algorithm with tweaked features and options.
It does 2 main functions :
    -Coding : Recieves an image and text as input (extracted from a file or written) ,creates a new image almost identical to the first containing the text 
              and provides a private key to decode .
    -Decoding : Recieves and image and key as input and creates a file containing the hidden text .

Supported formats :
    -Image : png , bitmap , gif .
    -text : txt, docx .
"""


# __Importing the python GUI library TKINTER__
from tkinter import *
from tkinter import filedialog
from tkinter import ttk
import tkinter.messagebox
from tkinter import simpledialog

# __Importing the python os,time and image library__
import os  # In order to get file size
import time  # In order to calculate execution time
import imageio as im #in order to read image dimensions

# __Importing the implemented Steganography algorithm__
import LSB
import LSBvGIF as LSBg

# __importing the module that would allow us to use multiple threads__
import threading

# __Importing the module that will allow us to read docx files__
import docx

#--------------------------->  GUI Class  <-----------------------------------#
#-----------------------------------------------------------------------------#
class SGUI():
    def __init__(self, master):
        '''Setting variables ,buttons , labels and tabs...'''
        self.master=master
        master.grid_rowconfigure(0, weight=1)
        master.grid_columnconfigure(0, weight=1)
        
        # __main_frame__      
        self.main_frame= ttk.Frame(master)
        self.main_frame.grid(row=0, column=0, sticky="nsew")
        
        # __Creating the encoding ,decoding and about tabs__
        self.tab_control = ttk.Notebook(self.main_frame)    
        self.encode_tab = ttk.Frame(self.tab_control)
        self.decode_tab = ttk.Frame(self.tab_control)
        self.encode_message_tab = ttk.Frame(self.tab_control)
        self.INFO_tab = ttk.Frame(self.tab_control)
        self.ABOUT_tab = ttk.Frame(self.tab_control)

        # __Adding the encoding ,decoding and about tabs__
        self.tab_control.add(self.encode_tab, text="Encode File")
        self.tab_control.add(self.encode_message_tab, text="Encode Message")
        self.tab_control.add(self.decode_tab, text="Decode Image")
        self.tab_control.add(self.INFO_tab, text="Info.")
        self.tab_control.add(self.ABOUT_tab, text="About")
        
        # __Variables__
        # Stores the path of the chosen file
        self.d_file_path = ""
        self.e_file_path = ""
        # Storing the path of the chosen image
        self.em_image_path = ""
        self.d_image_path = ""
        self.e_image_path = ""
        # Storing the placeholder content
        self.em_image = StringVar()
        self.d_image = StringVar()
        self.e_image = StringVar()
        self.d_file = StringVar()
        self.e_file = StringVar()
        self.em = StringVar()
        self.l0_text=StringVar()
        self.l1_text=StringVar()
        self.l2_text=StringVar()
        # Storing the size placeholder of each file
        self.e_image_size = StringVar()
        self.em_image_size = StringVar()
        self.e_file_size = StringVar()
        # Image & File sizes
        self.fsize = 0
        self.isize = 0
        # Control variables
        self.KEY = -1
        # The reset parameter & key_manager parameter
        self.k = -1
        self.x = -1
        # __Buttons__
        # -->Encoding Tab Buttons<--
                 # File selection button
        self.e_file_Button = ttk.Button(self.encode_tab, text="...", command=self.e_file_find_path)
                 # Image selection button
        self.e_image_Button = ttk.Button(self.encode_tab, text="...", command=self.e_image_find_path)
                 # Executing button
        self.e_go_Button = ttk.Button(self.encode_tab, text=" E N C O D E  ", command=self.thread_manager_e,state=DISABLED) 
                 # Adding buttons to the frame and managing their positions & style
        self.e_file_Button.place(x=440, y=70, width=30, height=20)
        self.e_image_Button.place(x=440, y=180, width=30, height=20)
        self.e_go_Button.place(x=1, y=225, width=475, height=70)

        # Main label
        self.l0 = Label(self.encode_tab, textvariable=self.l0_text)
        self.l0.place(x=1, y=1, relwidth=1 ) 
        self.l0_text.set(" ----->Select File and Image<----- ")  
        
        # File selection label & entry field
        self.e_file_label = Label(self.encode_tab, text=" Select a file :")
        self.e_file_label.place(x=1, y=40)
        self.e_file_entry = ttk.Entry(self.encode_tab, state=DISABLED, width=70, textvariable=self.e_file)
        self.e_file_entry.place(x=3, y=70)

        # Image selection label & entry field
        self.e_image_label = Label(self.encode_tab, text=" Select an image to encode :")
        self.e_image_label.place(x=1, y=150)
        self.e_image_entry = ttk.Entry(self.encode_tab, state=DISABLED, width=70, textvariable=self.e_image)
        self.e_image_entry.place(x=3, y=180)

        # File & Image size indicators
        self.e_file_size_label = Label(self.encode_tab, font=("Helvetica", 7), textvariable=self.e_file_size)
        self.e_file_size_label.place(x=1, y=92)
        self.e_image_size_label = Label(self.encode_tab, font=("Helvetica", 7), textvariable=self.e_image_size)
        self.e_image_size_label.place(x=1, y=202)

        # -->Encoding Message Tab Buttons<--
        # Main label
        self.l1 = Label(self.encode_message_tab,textvariable=self.l1_text)
        self.l1.place(x=1, y=1, relwidth=1)
        self.l1_text.set(" ----->Select Image and write your message<----- ")
        
        # image selection label & entry field       
        self.scrollbar = Scrollbar(self.encode_message_tab)
        self.scrollbar.place(x=433,y=43,height=100)
        self.TEXT = Text(self.encode_message_tab,wrap=WORD,yscrollcommand = self.scrollbar.set)
        self.TEXT.place(x=3, y=43, height=100, width=430)
        self.scrollbar.config( command = self.TEXT.yview )

        self.e_message_Button = ttk.Button(self.encode_message_tab, text=" E N C O D E  ", command=self.thread_manager_em ,state=DISABLED)
        self.e_message_Button.place(x=1, y=225, width=475, height=70)

        self.em_image_label = Label(self.encode_message_tab, text=" Select an image to encode:")
        self.em_image_label.place(x=1, y=150)

        self.em_txt_label = Label(self.encode_message_tab, text=" Place your text here :")
        self.em_txt_label.place(x=1, y=16)

        self.em_image_Button = ttk.Button(self.encode_message_tab, text="...", command=self.em_image_find_path)
        self.em_image_Button.place(x=440, y=180, width=30, height=20)

        self.em_image_entry = ttk.Entry(self.encode_message_tab, state=DISABLED, width=70, textvariable=self.em_image)
        self.em_image_entry.place(x=3, y=180)

        # -->Decoding Tab Buttons<--
        # File selection button
        self.d_file_Button = ttk.Button(self.decode_tab, text="...", command=self.d_file_find_path)
        # Image selection button
        self.d_image_Button = ttk.Button(self.decode_tab, text="...", command=self.d_image_find_path)
        # Executing button
        self.d_go_Button = ttk.Button(self.decode_tab, text=" D E C O D E   ", command=self.thread_manager_d,state=DISABLED)
        # Adding buttons to the frame and managing their positions & style
        self.d_file_Button.place(x=440, y=70, width=30, height=20)
        self.d_image_Button.place(x=440, y=180, width=30, height=20)

        self.d_go_Button.place(x=1, y=225, width=475, height=70)
        self.l2 = Label(self.decode_tab,textvariable=self.l2_text)
        self.l2.place(x=1, y=1, relwidth=1)
        self.l2_text.set(" ----->Select Image<----- ")
               
        # Labels & entry fields
        self.d_file_label = Label(self.decode_tab, text=" Select a destination file (optional) :")
        self.d_file_label.place(x=1, y=40)
        self.d_file_entry = ttk.Entry(self.decode_tab, state=DISABLED, width=70, textvariable=self.d_file)
        self.d_file_entry.place(x=3, y=70)

        self.d_image_label = Label(self.decode_tab, text=" Select an image to decode:")
        self.d_image_label.place(x=1, y=150)
        self.d_image_entry = ttk.Entry(self.decode_tab, state=DISABLED, width=70, textvariable=self.d_image)
        self.d_image_entry.place(x=3, y=180)

        # __INFO. tab__
        self.how = Label(self.INFO_tab, text="Wondering how to use SGUI by I&K ?")
        self.how.place(x=128, y=1)

        self.enc = Label(self.INFO_tab, text="Encode:  ")
        self.enc.place(x=125, y=45)

        self.encsel = Label(self.INFO_tab,
                            text="   •Select the text file (or input a message) \nand the image you wish to encode")
        self.encsel.place(x=140, y=70)

        self.pressenc = Label(self.INFO_tab, text="   •Press ENCODE")
        self.pressenc.place(x=140, y=105)

        self.dec = Label(self.INFO_tab, text="   Decode")
        self.dec.place(x=125, y=130)

        self.decsel = Label(self.INFO_tab, text="    •Select a destination file (if you want to)\n and the image you wish to decode")
        self.decsel.place(x=140, y=160)

        self.pressdec = Label(self.INFO_tab, text="   •Press DECODE")
        self.pressdec.place(x=140, y=200)

        self.author = Label(self.INFO_tab, text="AUTHORS : Ilyes Yahyaoui & Khaled Sellami")
        self.author.place(x=120, y=260)
        self.copyrigth = Label(self.INFO_tab, text="© 2018-2019 ALL RIGHTS RESERVED")
        self.copyrigth.place(x=140, y=278)

        # __Adding the tab control element to the frame__
        self.tab_control.pack(expand=1, fill=BOTH)

        # __ABOUT tab__
        self.info = Label(self.ABOUT_tab,text="SGUI by I&K is a steganography desktop application that uses \nthe modified least significant bit algorithm with tweaked features and options.")
        self.info.place(x=30, y=0)

        self.note = Label(self.ABOUT_tab, text="We would like to notify you that: ")
        self.note.place(x=125, y=45)

        self.utf = Label(self.ABOUT_tab, text="   • UTF-8 is supported.")
        self.utf.place(x=140, y=70)

        self.png = Label(self.ABOUT_tab, text="   • Png image format is supported.")
        self.png.place(x=140, y=100)

        self.bmp = Label(self.ABOUT_tab, text="   • Bmp image format is supported.")
        self.bmp.place(x=140, y=130)

        self.gif = Label(self.ABOUT_tab, text="   • Gif image format is supported.")
        self.gif.place(x=140, y=160)

        self.txt = Label(self.ABOUT_tab, text="   • Txt text format is supported.")
        self.txt.place(x=140, y=190)
        
        self.txt = Label(self.ABOUT_tab, text="   • Docx text format is supported.")
        self.txt.place(x=140, y=220)

        self.author = Label(self.ABOUT_tab, text="AUTHORS : Ilyes Yahyaoui & Khaled Sellami")
        self.author.place(x=120, y=260)
        self.copyrigth = Label(self.ABOUT_tab, text="© 2018-2019 ALL RIGHTS RESERVED")
        self.copyrigth.place(x=140, y=278)
        
        
        # __Setting up the progress bar frame__
        self.prog_bar_frame=ttk.Frame(master)
        self.prog_bar_frame.grid(row=0, column=0, sticky="nsew")
        self.bar = ttk.Progressbar(self.prog_bar_frame,length=260, mode="indeterminate")
        self.bar.place(x=20,y=40)
        self.l=Label(self.prog_bar_frame,text='Please wait ...')
        self.l.place(x=1,y=10,relwidth=1)
        
        
        self.main_frame.tkraise()
        

#--------------------->  File fetching function  <----------------------------#
        
    def e_file_find_path(self): 
        '''This function fetches is linked to the encode tab it fetches the path of the chosen file .
         File formats are limited to docx and txt.
        This function interacts dynamically with the user showing the file size and it's limits.'''    
        
        
        self.e_file_path = filedialog.askopenfilename(filetypes=(("Text file", "*.txt"), ("Word file", "*.docx"))) 
        self.e_file.set(self.e_file_path)
        if self.e_file_path != '':
            if  self.e_file_path[-3:]=="txt":
                self.fsize = os.path.getsize(self.e_file_path)
            else:
                doc = docx.Document(self.e_file_path)
                self.fsize = 0
                for para in doc.paragraphs:
                    self.fsize += len(para.text)
            self.e_file_size.set("size : " + str(self.fsize) + " bytes")

            if 0 < self.isize <= self.fsize*8:
                self.l0.config(fg="red")
                self.l0_text.set("--->Message can't be encoded in this Image<---\n--->Change Image or Message<---")
                self.e_go_Button.config(state=DISABLED)

            elif  32*self.fsize > self.isize >= self.fsize*8  :
                self.l0.config(fg="dark goldenrod")
                self.l0_text.set("--->Message can be encoded but riks of failing<---\n--->Choose a higher resolution image for a safe encoding<---")
                self.e_go_Button.config(state=NORMAL)
            else:
                self.l0.config(fg="green")
                self.l0_text.set("--->Message can be encoded safely<---")
                self.e_go_Button.config(state=NORMAL)
            if self.e_image_path=='':
                self.l0.config(fg="black")
                self.l0_text.set("--->Select Image<---")
                self.e_go_Button.config(state=DISABLED)
        else:
            self.l0.config(fg="black")
            if self.e_image_path=='':
                self.l0_text.set(" ----->Select File and Image<----- ")
            else :
                self.l0_text.set("--->Select File<---")
            self.e_go_Button.config(state=DISABLED)
            self.fsize = 0
            self.e_file_size.set("")
        return self.e_file_path
    

#--------------------->  File fetching function  <----------------------------#
   

    def d_file_find_path(self):
        '''This function fetches is linked to the decode tab it fetches the path of the chosen file .
        Choosing a decoding file is optional'''
        
        self.d_file_path = filedialog.askopenfilename(filetypes=(("Text file", "*.txt"),("Word file", "*.docx")))
        self.d_file.set(self.d_file_path)
        return self.d_file_path

#-------------------->  Image fetching function  <----------------------------#

    def e_image_find_path(self):
        '''This function fetches is linked to the encode tab it fetches the path of the chosen image.
        Image formats are limited to bmp ,png and gif.
        This function interacts dynamically with the user showing the image size and it's limits.'''
        
        self.e_image_path = filedialog.askopenfilename(
            filetypes=(("PNG", "*.png"), ("BITMAP", "*.bmp"), ("GIF", "*.gif")))
        self.e_image.set(self.e_image_path)
        if self.e_image_path != '':
            self.e_image.set(self.e_image_path)
            if self.e_image_path[-3:]=="gif" :
                imgl=im.mimread(self.e_image_path)
                img=imgl[0]
            else:
                img = im.imread(self.e_image_path)
            self.isize = img.shape[0] * img.shape[1] * 3
            self.e_image_size.set("This image can store at most " + str(self.isize//8) + " bytes")
            if 0 < self.isize <= self.fsize*8:
                self.l0.config(fg="red")
                self.l0_text.set("--->Message can't be encoded in this Image<---\n--->Change Image or Message<---")
                self.e_go_Button.config(state=DISABLED)
            elif  32*self.fsize > self.isize >= self.fsize*8  :
                self.l0.config(fg="dark goldenrod")
                self.l0_text.set("--->Message can be encoded but riks of failing<---\n--->Choose a higher resolution image for a safe encoding<---")
                self.e_go_Button.config(state=NORMAL)
            else:
                self.l0.config(fg="green")
                self.l0_text.set("--->Message can be encoded safely<---")
                self.e_go_Button.config(state=NORMAL)
            if self.fsize==0:
                self.l0.config(fg="black")
                self.l0_text.set("--->Select File<---")
                self.e_go_Button.config(state=DISABLED)
        else:
            self.l0.config(fg="black")
            if self.fsize==0:
                self.l0_text.set(" ----->Select File and Image<----- ")
            else:
                self.l0_text.set("--->Select Image<---")
            self.e_go_Button.config(state=DISABLED)
            self.isize = 0
            self.e_image_size.set("")
        return self.e_image_path


#-------------------->  Image fetching function  <----------------------------#

    def em_image_find_path(self):
        '''This function fetches is linked to the encode message tab it fetches the path of the chosen image.
        Image formats are limited to bmp ,png and gif.
        This function interacts dynamically with the user showing the image size and it's limits.'''  
        
        self.em_image_path = filedialog.askopenfilename(
            filetypes=(("PNG", "*.png"), ("BITMAP", "*.bmp"), ("GIF", "*.gif"))) 
        self.em_image.set(self.em_image_path)
        if self.em_image_path=="":
            self.l1.config(fg="black")
            self.l1_text.set("--->Select Image<---")
            self.e_message_Button.config(state=DISABLED)
        else:
            self.l1.config(fg="green")
            self.l1_text.set("--->Encoder is ready<---")
            self.e_message_Button.config(state=NORMAL)
            
            if self.em_image_path[-3:]=="gif" :
                imgl=im.mimread(self.em_image_path)
                img=imgl[0]
            else:
                img = im.imread(self.em_image_path)
            self.imsize = img.shape[0] * img.shape[1] * 3
        return self.em_image_path
    
#-------------------->  Image fetching function  <----------------------------#

    def d_image_find_path(self):
        '''This function fetches is linked to the decode tab it fetches the path of the chosen image.
        Image formats are limited to bmp ,png and gif.
        This function interacts dynamically with the user showing the image size and it's limits.''' 
        
        self.d_image_path = filedialog.askopenfilename(
            filetypes=(("PNG", "*.png"), ("BITMAP", "*.bmp"), ("GIF", "*.gif")))
        self.d_image.set(self.d_image_path)
        if self.d_image_path=="":
            self.l2.config(fg="black")
            self.l2_text.set("--->Select Image<---")
            self.d_go_Button.config(state=DISABLED)
        else:
            self.l2.config(fg="green")
            self.l2_text.set("--->Decoder is ready<---")
            self.d_go_Button.config(state=NORMAL)
        return self.d_image_path

#-------------------->  Reset parameters function  <--------------------------#


    def reset(self):
        '''This function resets the placeholders, size labels and the path variables__
        x=0 reset paramters of the Encode tab /x=1 reset paramters of the Decode tab/
        x=# reset paramters of the Encode message tab'''
        
        if self.x == 0:
            self.e_file_path = ""
            self.e_image_path = ""
            self.e_file.set(self.e_image_path)
            self.e_image.set(self.e_image_path)
            self.e_image_size.set("")
            self.e_file_size.set("")
            self.isize = 0
            self.fsize = 0
            self.KEY = -1
            self.l0.config(fg="black")
            self.l0_text.set(" ----->Select File and Image<----- ")
            self.e_go_Button.config(state=DISABLED)

        elif self.x == 1:
            self.d_file_path = ""
            self.d_image_path = ""
            self.d_file.set(self.d_file_path)
            self.d_image.set(self.d_image_path)
            self.l2.config(fg="black")
            self.l2_text.set("--->Select Image<---")
            self.d_go_Button.config(state=DISABLED)
            self.isize = 0
            self.fsize = 0
            self.KEY = -1

        else:
            self.em_image_path = ""
            self.em_image.set(self.em_image_path)
            self.em.set("")
            self.l1.config(fg="black")
            self.l1_text.set("--->Select Image<---")
            self.e_message_Button.config(state=DISABLED)       
            self.TEXT.place(x=3, y=43, height=100, width=430)
            self.scrollbar.place(x=433,y=43,height=100)
            
#---------------------->  Key managing function  <----------------------------#

    def key_manager(self):
        '''This function is responsable for providing the Encryption key and outputting it
        k=0 show the key/ k=# demand key input'''       
    
        if self.k == 0:
            keywindow = Toplevel()
            keywindow.iconbitmap('icons/key.ico')
            keywindow.title("Decoding Key")
            keywindow.resizable(width=False, height=False)
            w = Text(keywindow, height=1, width=30)
            w.insert(1.0, "KEY : " + str(self.KEY))
            w.pack()
            w.configure(bg=keywindow.cget('bg'), relief=FLAT, state=DISABLED)

        else:
            self.KEY = tkinter.simpledialog.askinteger("Decode Key", "Set your KEY")
            if self.KEY==None:
                self.reset()
                
#---------------------->  Threading functions  <-------------------------------#
   
    def thread_manager_e(self):
        self.prog_bar_frame.tkraise()
        ws = self.master.winfo_screenwidth()
        hs = self.master.winfo_screenheight() 
        w = 300  
        h = 80  
        x = (ws / 2) - (w / 2)
        y = (hs / 2) - (h / 2)
        self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))
        self.bar.start(10)
        if self.e_image_path[-3:]=="gif": 
            self.encoder()
        else:
            T1=threading.Thread(target=self.encoder)
            T1.start()
            
            self.l0.config(fg="green")
            self.l0_text.set("--->DONE<---")
        return
    
    def thread_manager_em(self):  
        self.prog_bar_frame.tkraise()
        ws = self.master.winfo_screenwidth()
        hs = self.master.winfo_screenheight() 
        w = 300  
        h = 80  
        x = (ws / 2) - (w / 2)
        y = (hs / 2) - (h / 2)
        self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))
        self.bar.start(10)
        if self.em_image_path[-3:]=="gif": 
            self.encoder_msg()
            
        else:
            T2=threading.Thread(target=self.encoder_msg)
            T2.start()
            
            self.l1.config(fg="green")
            self.l1_text.set("--->DONE<---")
        
    def thread_manager_d(self):
        self.k = 1
        self.x = 1
        self.key_manager()
        if self.KEY==-1:
            return
        elif self.KEY<16777216 :
            tkinter.messagebox.showinfo("ERROR !","Wrong key .")
            return
        self.prog_bar_frame.tkraise()
        ws = self.master.winfo_screenwidth()
        hs = self.master.winfo_screenheight() 
        w = 300  
        h = 80  
        x = (ws / 2) - (w / 2)
        y = (hs / 2) - (h / 2)
        self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))
        self.bar.start(10)
        T1=threading.Thread(target=self.decoder)
        T1.start()
        self.l2.config(fg="green")
        self.l2_text.set("--->DONE<---")
            
        
            
#-------------------->  Encode message function  <----------------------------#

    # __Encoding functions__
    def encoder_msg(self): 
        '''LSB encoding for text'''
        self.k = 0
        self.x = 3
        M = self.TEXT.get("1.0","end-1c")
        if self.em_image_path != "" and M == "":
            #restore previous size and frame
            self.bar.stop()
            ws = self.master.winfo_screenwidth()  
            hs = self.master.winfo_screenheight()  
            w = 480  
            h = 320  
            x = (ws / 2) - (w / 2)
            y = (hs / 2) - (h / 2)
            self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))        
            self.main_frame.tkraise()
            tkinter.messagebox.showerror('ERROR !', "No text set")

        else:
            if len(str(M))*8>self.imsize :
                #restore previous size and frame
                self.bar.stop()
                ws = self.master.winfo_screenwidth()  
                hs = self.master.winfo_screenheight()  
                w = 480  
                h = 320  
                x = (ws / 2) - (w / 2)
                y = (hs / 2) - (h / 2)
                self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))        
                self.main_frame.tkraise()
                tkinter.messagebox.showinfo("ERROR !",
                                            "This message can't be encoded in this image.\nYou should choose a shorter message or a higher resolution image.")
                return
            start = time.time()
            if self.em_image_path[-3:]=="gif":
                TUP = LSBg.codeGIF(self.em_image_path,str(M))
            else: 
                TUP = LSB.code(self.em_image_path,str(M))
            self.KEY = TUP[0]
            xt = str(time.time() - start)
            t = xt[:3]
            
            
            #restore previous size and frame
            self.bar.stop()
            ws = self.master.winfo_screenwidth()  
            hs = self.master.winfo_screenheight()  
            w = 480  
            h = 320  
            x = (ws / 2) - (w / 2)
            y = (hs / 2) - (h / 2)
            self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))        
            self.main_frame.tkraise()
            
            if self.KEY==-1 :
                tkinter.messagebox.showinfo("ERROR !",
                                            "This message can't be encoded in this image.\nYou should choose a shorter message or a higher resolution image.")
            else:
                self.key_manager()
                path = self.em_image_path[:self.em_image_path.rfind('/')]
                A = tkinter.messagebox.askquestion('Done !', "Encoding complete in " + t + " seconds.\nYou will find your encoded image at  " + path + "\nWould you like to display the encoded image and the original one ?")
                if A == 'yes':
                    if self.em_image_path[-3:]=='gif':
                        LSBg.plot(self.em_image_path, TUP[1], path)
                    else:
                        LSB.plot(self.em_image_path, TUP[1])
            self.reset()
            return
        

#---------------------->  Encoding function  <--------------------------------#     
    def encoder(self):
        '''LSB encoding for files'''
        self.k = 0
        self.x = 0
        if self.e_file_path[-3:] != "txt" : 
            doc = docx.Document(self.e_file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + '\n'
        else:
            try:
                with open(self.e_file_path) as file:
                    text=file.read()
            except UnicodeDecodeError :
                with open(self.e_file_path,encoding="utf-8") as file:
                    text=file.read()
        path = self.e_image_path[:self.e_image_path.rfind('/')]       
        start = time.time()
        if self.e_image_path[-3:]=="gif":            
            TUP = LSBg.codeGIF(self.e_image_path, text)
        else :
            TUP = LSB.code(self.e_image_path, text)
        self.KEY = TUP[0]        
        x = str(time.time() - start)
        t = x[:3]
        
        
        #restore previous size and frame
        self.bar.stop()
        ws = self.master.winfo_screenwidth()  
        hs = self.master.winfo_screenheight()  
        w = 480  
        h = 320  
        x = (ws / 2) - (w / 2)
        y = (hs / 2) - (h / 2)
        self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))        
        self.main_frame.tkraise()
        
        if self.KEY== -1 : 
            tkinter.messagebox.showinfo("ERROR !",
                                        "This message can't be encoded in this image.\nYou should choose a shorter lessage or a higher resolution image.")
            return
        else:
            self.key_manager()
            A = tkinter.messagebox.askquestion('Done !', "Encoding complete in " + t + " seconds.\nYou will find your encoded image at  " + path + "\nWould you like to display the encoded image and the original one ?")
            if A == 'yes':
                if self.e_image_path[-3:]=='gif':
                    LSBg.plot(self.e_image_path, TUP[1], path)
                else:
                    LSB.plot(self.e_image_path, TUP[1])
        
        self.reset()  
        return

#------------------------->  Decoding function  <-----------------------------#
        
    def decoder(self):
        '''LSB decoding'''
 
        start = time.time()
        if self.d_image_path[-3:]=="gif":
            path=LSBg.decodeGIF(self.d_image_path, self.d_file_path, self.KEY)
        else:
            path=LSB.decode(self.d_image_path, self.d_file_path, self.KEY)
        
        
        #restore previous size and frame
        self.bar.stop()
        ws = self.master.winfo_screenwidth()  
        hs = self.master.winfo_screenheight()  
        w = 480  
        h = 320  
        x = (ws / 2) - (w / 2)
        y = (hs / 2) - (h / 2)
        self.master.geometry('%dx%d+%d+%d' % (w, h, x, y))        
        self.main_frame.tkraise()
        
        
        if path=="ERROR : wrong key": 
            tkinter.messagebox.showinfo("ERROR !","Wrong key or wrong image.")
            self.reset()
            return
        x = str(time.time() - start)
        t = x[:3]
        A=tkinter.messagebox.askquestion("Done !",
                                        "Decoding complete in " + t + " seconds.\n You will find your hidden text at "
                                        + path+"\nWould you like to display the decoded text ?")
        if A=='yes':
            if path[-3:] != "txt" :
                doc = docx.Document(path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + '\n'
            else:
                with open(path,encoding="utf-8") as file:
                    text=file.read()
            window = Toplevel()
            window.title("Message")
            window.iconbitmap('icons/display.ico')
            scrollbar = Scrollbar(window)
            scrollbar.pack( side = RIGHT, fill = Y )
            w = Text(window,wrap=WORD,yscrollcommand = scrollbar.set)
            w.insert(INSERT,text)
            w.pack(fill=BOTH, expand=TRUE)
            w.configure(bg=window.cget('bg'), relief=FLAT,state=DISABLED)
            scrollbar.config( command = w.yview )
        self.reset()
        return
    

#------------------------->  End of Class  <----------------------------------#
#-----------------------------------------------------------------------------#


# __Mainwindow's parameters__
mainwindow = Tk()
Object = SGUI(mainwindow)
mainwindow.title("SGUI by I&K")
mainwindow.resizable(width=False, height=False)  # Fixed height and width
# Making the window pop in the middle of the screen
ws = mainwindow.winfo_screenwidth()  # Width of the screen
hs = mainwindow.winfo_screenheight()  # Height of the screen
w = 480  # Width of the mainwindow
h = 320  # Height of the mainwindow

# Calculate x and y coordinates for the Tk Mainwindow
x = (ws / 2) - (w / 2)
y = (hs / 2) - (h / 2)
# Setting the dimensions of the screen and where it is placed
mainwindow.geometry('%dx%d+%d+%d' % (w, h, x, y))

mainwindow.iconbitmap("icons/main.ico")

mainwindow.mainloop()
